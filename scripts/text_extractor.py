import os, fitz, pandas as pd
from docx import Document
from bs4 import BeautifulSoup

RAW_DIR = '../raw_documents'
TEXT_DIR = '../extracted_text'
LOG_PATH = '../logs/parse_errors.log'

def ensure_dirs():
    """Ensure required directories exist"""
    for dir_path in [RAW_DIR, TEXT_DIR, os.path.dirname(LOG_PATH)]:
        os.makedirs(dir_path, exist_ok=True)

def extract_pdf(path):
    """Extract text from PDF file"""
    print(f"  Extracting PDF: {os.path.basename(path)}")
    doc = fitz.open(path)
    text_content = []
    for page_num, page in enumerate(doc, 1):
        page_text = page.get_text()
        if page_text.strip():
            text_content.append(f"=== PAGE {page_num} ===\n{page_text}")
    doc.close()
    return "\n\n".join(text_content)

def extract_excel(path):
    """Extract text from Excel file"""
    print(f"  Extracting Excel: {os.path.basename(path)}")
    try:
        # Try to read all sheets
        excel_file = pd.ExcelFile(path)
        sheet_contents = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name)
            sheet_text = f"=== SHEET: {sheet_name} ===\n{df.to_string(index=False)}"
            sheet_contents.append(sheet_text)
        
        return "\n\n".join(sheet_contents)
    except Exception as e:
        # Fallback: try reading as CSV-like
        df = pd.read_excel(path)
        return df.to_string(index=False)

def extract_docx(path):
    """Extract text from DOCX file"""
    print(f"  Extracting DOCX: {os.path.basename(path)}")
    doc = Document(path)
    
    content = []
    
    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        content.append("=== DOCUMENT TEXT ===")
        content.extend(paragraphs)
    
    # Extract tables
    if doc.tables:
        content.append("\n=== TABLES ===")
        for i, table in enumerate(doc.tables, 1):
            content.append(f"\n--- Table {i} ---")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    content.append(row_text)
    
    return "\n".join(content)

def extract_html(path):
    """Extract text from HTML file"""
    print(f"  Extracting HTML: {os.path.basename(path)}")
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        soup = BeautifulSoup(f, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and clean it up
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text

def extract_text(file_path):
    """Extract text from various file formats"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return extract_pdf(file_path)
        elif file_ext in ['.xls', '.xlsx']:
            return extract_excel(file_path)
        elif file_ext == '.docx':
            return extract_docx(file_path)
        elif file_ext in ['.html', '.htm']:
            return extract_html(file_path)
        else:
            print(f"  Unsupported file type: {file_ext}")
            return f"Unsupported file type: {file_ext}"
            
    except Exception as e:
        error_msg = f"ERROR: {file_path} - {str(e)}\n"
        print(f"  {error_msg.strip()}")
        with open(LOG_PATH, 'a') as log:
            log.write(error_msg)
        return f"Error extracting text: {str(e)}"

def main():
    """Main function to extract text from all downloaded documents"""
    print("Text Extractor Starting...")
    print("=" * 50)
    
    ensure_dirs()
    
    if not os.path.exists(RAW_DIR):
        print(f"Error: Raw documents directory not found: {RAW_DIR}")
        return
    
    raw_files = [f for f in os.listdir(RAW_DIR) if os.path.isfile(os.path.join(RAW_DIR, f))]
    
    if not raw_files:
        print(f"No files found in {RAW_DIR}")
        print("Run rss_harvester.py first to download documents")
        return
    
    print(f"Found {len(raw_files)} files to process")
    
    processed = 0
    skipped = 0
    
    for fname in raw_files:
        print(f"\nProcessing: {fname}")
        
        fpath = os.path.join(RAW_DIR, fname)
        base_name = os.path.splitext(fname)[0]
        tpath = os.path.join(TEXT_DIR, f"{base_name}.txt")
        
        if os.path.exists(tpath):
            print(f"  Text file already exists, skipping: {base_name}.txt")
            skipped += 1
            continue
        
        text_content = extract_text(fpath)
        
        if text_content:
            try:
                with open(tpath, 'w', encoding='utf-8') as out:
                    # Add metadata header
                    metadata = f"""=== EXTRACTED TEXT METADATA ===
Source File: {fname}
Source Path: {fpath}
Extraction Date: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}
File Size: {os.path.getsize(fpath)} bytes
=== END METADATA ===

"""
                    out.write(metadata + text_content)
                
                print(f"  Saved extracted text: {base_name}.txt")
                processed += 1
                
            except Exception as e:
                error_msg = f"ERROR saving {tpath}: {str(e)}\n"
                print(f"  {error_msg.strip()}")
                with open(LOG_PATH, 'a') as log:
                    log.write(error_msg)
        else:
            print(f"  No text content extracted from {fname}")
    
    print("\n" + "=" * 50)
    print(f"Text extraction complete!")
    print(f"Processed: {processed} files")
    print(f"Skipped: {skipped} files")
    print(f"Extracted text saved to: {TEXT_DIR}")
    print(f"Check {LOG_PATH} for any errors")

if __name__ == "__main__":
    main()